import os
import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

import PyPDF2
from docx import Document
import openpyxl
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

@dataclass
class ProcessedDocument:
    content: str
    metadata: Dict[str, Any]
    source: str
    doc_type: str

class DocumentProcessor:
    """Procesador de documentos con soporte para múltiples formatos"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        self.logger = logging.getLogger(__name__)
    
    def process_document(self, file_path: str) -> List[ProcessedDocument]:
        """Procesa un documento según su extensión"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"El archivo {file_path} no existe")
        
        extension = file_path.suffix.lower()
        
        processors = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.md': self._process_markdown,
            '.json': self._process_json,
            '.xlsx': self._process_xlsx,
            '.xls': self._process_xlsx
        }
        
        if extension not in processors:
            raise ValueError(f"Formato {extension} no soportado")
        
        try:
            content = processors[extension](file_path)
            return self._chunk_content(content, str(file_path), extension[1:])
        except Exception as e:
            self.logger.error(f"Error procesando {file_path}: {str(e)}")
            raise
    
    def _process_pdf(self, file_path: Path) -> str:
        """Procesa archivos PDF"""
        content = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
        return content.strip()
    
    def _process_docx(self, file_path: Path) -> str:
        """Procesa archivos DOCX"""
        doc = Document(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        return content.strip()
    
    def _process_txt(self, file_path: Path) -> str:
        """Procesa archivos TXT"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    
    def _process_markdown(self, file_path: Path) -> str:
        """Procesa archivos Markdown"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
            # Convertir a texto plano manteniendo estructura
            html = markdown.markdown(md_content)
            # Para simplicidad, devolver el markdown original
            return md_content.strip()
    
    def _process_json(self, file_path: Path) -> str:
        """Procesa archivos JSON"""
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            # Convertir JSON a texto legible
            return json.dumps(data, indent=2, ensure_ascii=False)
    
    def _process_xlsx(self, file_path: Path) -> str:
        """Procesa archivos Excel"""
        workbook = openpyxl.load_workbook(file_path)
        content = ""
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            content += f"### Hoja: {sheet_name}\n\n"
            
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                content += " | ".join(row_data) + "\n"
            content += "\n"
        
        return content.strip()
    
    def _chunk_content(self, content: str, source: str, doc_type: str) -> List[ProcessedDocument]:
        """Divide el contenido en chunks"""
        if not content.strip():
            return []
        
        chunks = self.text_splitter.split_text(content)
        processed_docs = []
        
        for i, chunk in enumerate(chunks):
            metadata = {
                'source': source,
                'doc_type': doc_type,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            
            processed_docs.append(
                ProcessedDocument(
                    content=chunk,
                    metadata=metadata,
                    source=source,
                    doc_type=doc_type
                )
            )
        
        return processed_docs



